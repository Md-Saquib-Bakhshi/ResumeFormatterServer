import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE # Correct import, already there
from typing import Dict, Any, List

def generate_docx_from_json(parsed_data: Dict[str, Any]) -> bytes:
    """
    Generates a DOCX resume from structured JSON data.
    """
    document = Document()

    # Define some basic styles
    styles = document.styles
    
    # Ensure 'Normal' style font size
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)

    # Heading 1 for Name
    if 'Heading 1' not in styles:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.size = Pt(24)
        h1_style.font.bold = True
    # If it exists, we don't need to do anything, it will be used by name_para

    # Heading 2 for sections
    if 'Heading 2' not in styles:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.size = Pt(14)
        h2_style.font.bold = True
    # If it exists, use it

    # Check for 'List Bullet' style or define a fallback
    # Most Word documents have 'List Bullet' by default.
    # We'll use this for bullet points.
    list_bullet_style = None
    if 'List Bullet' in styles:
        list_bullet_style = 'List Bullet'
    else:
        # Fallback: if 'List Bullet' isn't there, create a simple bullet style
        # This is less common but good for robustness.
        try:
            list_bullet_style = styles.add_style('MyBulletList', WD_STYLE_TYPE.PARAGRAPH)
            list_bullet_style.base_style = styles['List Paragraph'] # Inherit from List Paragraph
            list_bullet_style.paragraph_format.left_indent = Inches(0.5)
            # You might need to manually set font/bullet char for 'MyBulletList' for full control
            # but usually just setting base_style to 'List Paragraph' is enough for basic bullets.
            # However, for true bullets, associating with a numbering format is best.
            # For this simple example, we'll rely on built-in 'List Bullet' and fallback to manual.
            list_bullet_style = 'MyBulletList' # Use our custom style
            print("Custom bullet list style created.") # For debugging
        except Exception as e:
            print(f"Could not create custom bullet style, falling back to manual: {e}")
            list_bullet_style = None # Indicate fallback to manual bullet in code

    # Add Basic Details
    basic_details = parsed_data.get("basic_details", {})
    name = basic_details.get("name", "N/A")
    email = basic_details.get("email", "N/A")
    phone = basic_details.get("phone", "N/A")
    links = basic_details.get("links", {})

    # Name
    name_para = document.add_paragraph(name, style='Heading 1')
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact_info = []
    if email != "N/A":
        contact_info.append(email)
    if phone != "N/A":
        contact_info.append(phone)
    for link_name, url in links.items():
        if url and url != "N/A":
            contact_info.append(f"{link_name.capitalize()}: {url}")
    
    if contact_info:
        contact_para = document.add_paragraph(" | ".join(contact_info), style='Normal')
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph("") # Add a blank line for spacing

    # Professional Summary
    professional_summary = parsed_data.get("professional_summary", "N/A")
    if professional_summary != "N/A" and professional_summary.strip():
        document.add_paragraph("Professional Summary", style='Heading 2')
        document.add_paragraph(professional_summary, style='Normal')
        document.add_paragraph("")

    # Technical Expertise
    technical_expertise = parsed_data.get("technical_expertise", [])
    if technical_expertise:
        document.add_paragraph("Technical Expertise", style='Heading 2')
        # Join skills with commas, potentially splitting into multiple lines if too long
        skills_text = ", ".join(technical_expertise)
        document.add_paragraph(skills_text, style='Normal')
        document.add_paragraph("")

    # Professional Experience
    professional_experience = parsed_data.get("professional_experience", [])
    if professional_experience:
        document.add_paragraph("Professional Experience", style='Heading 2')
        for exp in professional_experience:
            company = exp.get("company", "N/A")
            role = exp.get("role", "N/A")
            date_range = exp.get("date_range", "N/A")
            client_engagement = exp.get("client_engagement", "N/A")
            program = exp.get("program", "N/A")
            responsibilities = exp.get("responsibilities", [])

            exp_para = document.add_paragraph(f"{role}", style='Normal')
            exp_para.add_run(f" at {company}").bold = True
            exp_para.add_run(f" ({date_range})")

            if client_engagement and client_engagement != "N/A":
                document.add_paragraph(f"Client Engagement: {client_engagement}", style='Normal')
            if program and program != "N/A":
                document.add_paragraph(f"Program: {program}", style='Normal')

            if responsibilities:
                for resp in responsibilities:
                    if resp and resp != "N/A":
                        if list_bullet_style: # Use the determined bullet style
                             document.add_paragraph(resp, style=list_bullet_style)
                        else: # Fallback to manual bullet if no proper list style found/created
                             document.add_paragraph(f"• {resp}", style='Normal')
            document.add_paragraph("") # Space between experiences

    # Certifications
    certifications = parsed_data.get("certifications", [])
    if certifications:
        document.add_paragraph("Certifications", style='Heading 2')
        for cert in certifications:
            title = cert.get("title", "N/A")
            date = cert.get("date", "N/A")
            if title != "N/A":
                if list_bullet_style: # Use the determined bullet style
                    cert_para = document.add_paragraph(title, style=list_bullet_style)
                else: # Fallback to manual bullet
                    cert_para = document.add_paragraph(f"• {title}", style='Normal')
                if date and date != "N/A":
                    cert_para.add_run(f" ({date})")
        document.add_paragraph("")

    # Save the document to an in-memory byte stream
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()