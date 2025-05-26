import io
import json
import asyncio
import concurrent.futures
import re
import os
import logging
import zipfile
from typing import List, Dict, Any, Tuple, Optional

from PIL import Image
import numpy as np
import fitz # PyMuPDF
import easyocr
from docx import Document # For .docx files
from pydocx import PyDocX # For .doc files
import aiofiles # For async file I/O
import shutil # For removing directories

# Import the configured LLM client and prompt from utils
from .llm_config import azure_openai_client, AZURE_DEPLOYMENT_NAME, LLM_RESUME_PARSING_PROMPT
# Update import for OutputFormat
from ...utils.job_manager import job_manager, JobStatus, OutputFormat 

# Import your PDF generator
from .generators.pdf_generator import generate_pdf_from_json
# Import your DOCX generator
from .generators.docx_generator import generate_docx_from_json 

logger = logging.getLogger(__name__) # Get a logger instance for this module

# Initialize EasyOCR reader once globally for performance.
reader = easyocr.Reader(['en'], gpu=False) # Set gpu=False if you don't have CUDA/GPU setup

# Initialize a global thread pool for running synchronous OCR and LLM calls.
# Increased max_workers to potentially handle more concurrent OCR/LLM calls.
ocr_llm_executor = concurrent.futures.ThreadPoolExecutor(max_workers=8)

# --- Define your project's download directory ---
# This path assumes your 'downloads' folder is at the 'src/server' level.
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..'))
DOWNLOAD_DIR = os.path.join(PROJECT_ROOT, 'downloads')

# Ensure the download directory exists when the service module is loaded
os.makedirs(DOWNLOAD_DIR, exist_ok=True)
logger.info(f"Ensured download directory exists at: {DOWNLOAD_DIR}")
# --- End of new additions ---


async def _call_llm_for_resume_parsing(resume_text: str) -> Tuple[Dict[str, Any], Dict[str, Any]]:
    """
    Internal helper to call Azure OpenAI for structured resume data extraction.
    This function is designed to be run within a thread pool executor.
    Returns parsed data and usage info.
    """
    if azure_openai_client is None:
        logger.error("Azure OpenAI client is not initialized. Cannot perform LLM call.")
        raise RuntimeError("Azure OpenAI client is not initialized. Please set up API credentials.")

    loop = asyncio.get_event_loop()
    
    prompt_messages = [
        {"role": "system", "content": LLM_RESUME_PARSING_PROMPT},
        {"role": "user", "content": f"Resume Text:\n\n{resume_text}\n\nExtract the structured resume data:"}
    ]

    llm_output_str = ""
    usage_info = {}
    try:
        response = await loop.run_in_executor(
            ocr_llm_executor,
            lambda: azure_openai_client.chat.completions.create(
                model=AZURE_DEPLOYMENT_NAME,
                messages=prompt_messages,
                response_format={"type": "json_object"},
                temperature=0.1,
                max_tokens=5000
            )
        )
        
        if response.usage:
            usage_info = {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            }
        else:
            logger.warning("LLM response did not contain usage information.")

        llm_output_str = response.choices[0].message.content
        llm_data = json.loads(llm_output_str)
        
        # --- Post-processing to ensure adherence to schema and apply defaults ---
        parsed_resume = {
            "basic_details": {
                "name": "N/A", "email": "N/A", "phone": "N/A",
                "links": {}
            },
            "technical_expertise": [],
            "certifications": [],
            "professional_summary": "N/A",
            "professional_experience": []
        }

        # Safely populate basic_details
        if "basic_details" in llm_data and isinstance(llm_data["basic_details"], dict):
            parsed_resume["basic_details"]["name"] = llm_data["basic_details"].get("name", "N/A")
            parsed_resume["basic_details"]["email"] = llm_data["basic_details"].get("email", "N/A")
            parsed_resume["basic_details"]["phone"] = llm_data["basic_details"].get("phone", "N/A")
            if "links" in llm_data["basic_details"] and isinstance(llm_data["basic_details"]["links"], dict):
                for link_key, link_val in llm_data["basic_details"]["links"].items():
                    if link_val and str(link_val).strip() != "N/A":
                        parsed_resume["basic_details"]["links"][link_key.lower()] = str(link_val).strip()

        # Safely populate technical_expertise
        if "technical_expertise" in llm_data and isinstance(llm_data["technical_expertise"], list):
            parsed_resume["technical_expertise"] = sorted(list(set([
                str(s).strip() for s in llm_data["technical_expertise"] if str(s).strip() and str(s).strip() != "N/A"
            ])))

        # Safely populate certifications
        if "certifications" in llm_data and isinstance(llm_data["certifications"], list):
            certs = []
            for cert_item in llm_data["certifications"]:
                if isinstance(cert_item, dict):
                    title = cert_item.get("title", "N/A")
                    date = cert_item.get("date", "N/A")
                    if title and title != "N/A":
                        certs.append({"title": str(title).strip(), "date": str(date).strip()})
                elif isinstance(cert_item, str) and str(cert_item).strip() != "N/A":
                    certs.append({"title": str(cert_item).strip(), "date": "N/A"})
            seen_titles = set()
            unique_certs = []
            for cert in certs:
                if cert["title"].lower() not in seen_titles:
                    unique_certs.append(cert)
                    seen_titles.add(cert["title"].lower())
            parsed_resume["certifications"] = sorted(unique_certs, key=lambda x: x["title"].lower())

        # Safely populate professional_summary
        if "professional_summary" in llm_data:
            parsed_resume["professional_summary"] = str(llm_data["professional_summary"]).strip() or "N/A"
        
        # Safely populate professional_experience
        experiences = []
        if "professional_experience" in llm_data and isinstance(llm_data["professional_experience"], list):
            for exp in llm_data["professional_experience"]:
                if isinstance(exp, dict):
                    responsibilities = exp.get("responsibilities", [])
                    if isinstance(responsibilities, str):
                        responsibilities = [responsibilities]
                    
                    experiences.append({
                        "company": exp.get("company", "N/A"),
                        "date_range": exp.get("date_range", "N/A"),
                        "role": exp.get("role", "N/A"),
                        "client_engagement": exp.get("client_engagement", "N/A"),
                        "program": exp.get("program", "N/A"),
                        "responsibilities": [str(r).strip() for r in responsibilities if str(r).strip() != "N/A"]
                    })
        
        parsed_resume["professional_experience"] = sorted(
            experiences, 
            key=lambda x: len(x.get("responsibilities", [])), 
            reverse=True
        )

        # Log the parsed JSON data
        logger.info(f"Parsed JSON output from LLM: {json.dumps(parsed_resume, indent=2)}")
        
        return parsed_resume, usage_info

    except json.JSONDecodeError as e:
        logger.error(f"LLM did not return valid JSON. Raw output: {llm_output_str if 'llm_output_str' in locals() else 'N/A'}. Error: {e}. Resume text snippet: {resume_text[:500]}", exc_info=True)
        raise ValueError(f"LLM response error: {e}")
    except Exception as e:
        logger.error(f"Failed to call Azure OpenAI or process its response: {e}", exc_info=True)
        raise ValueError(f"LLM call error: {e}")


def _extract_text_from_docx(doc_bytes: bytes) -> str:
    """
    Synchronously extracts plain text from a .docx file using python-docx.
    """
    text_content = []
    try:
        document = Document(io.BytesIO(doc_bytes))
        for para in document.paragraphs:
            text_content.append(para.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text_content.append(paragraph.text)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}", exc_info=True)
        raise ValueError(f"Could not read DOCX file: {e}")
    return "\n".join(text_content)


def _extract_text_from_doc(doc_bytes: bytes) -> str:
    """
    Synchronously extracts plain text from a .doc file using pydocx.
    """
    text_content = ""
    try:
        html_content = PyDocX.to_html(io.BytesIO(doc_bytes))
        cleanr = re.compile('<.*?>')
        text_content = re.sub(cleanr, '', html_content)
    except Exception as e:
        logger.error(f"Failed to extract text from DOC: {e}", exc_info=True)
        raise ValueError(f"Could not read DOC file: {e}")
    return text_content


async def process_single_resume_file(
    file_bytes: bytes, original_filename: str, content_type: str
) -> Dict[str, Any]:
    """
    Processes a single resume file (PDF, DOCX, DOC) to extract structured data.
    This is now a helper for batch processing.
    """
    all_extracted_text_segments = []

    if content_type == "application/pdf":
        logger.info(f"Starting PDF text extraction and OCR for '{original_filename}'.")
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        loop = asyncio.get_event_loop()
        ocr_tasks = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            text_from_pdf_layer = page.get_text("text")
            if text_from_pdf_layer.strip():
                all_extracted_text_segments.append(f"--- Page {page_num + 1} (PDF Text Layer) ---\n")
                all_extracted_text_segments.append(text_from_pdf_layer)
            
            perform_full_page_ocr = False
            native_text_len = len(text_from_pdf_layer.strip())

            # Heuristic for determining if full page OCR is needed
            if native_text_len < 100: # Very little native text
                perform_full_page_ocr = True
            else:
                # Check for common sections if there's some text but might be incomplete
                common_section_keywords = ["experience", "education", "skills", "certifications", "summary", "projects", "awards", "licenses", "certified"]
                if not any(keyword in text_from_pdf_layer.lower() for keyword in common_section_keywords):
                    perform_full_page_ocr = True 

            image_list = page.get_images(full=True)
            for img_index, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = doc.extract_image(xref)
                img_bytes = base_image.get("image")
                img_ext = base_image.get("ext")

                if not img_bytes:
                    logger.warning(f"No image bytes extracted for embedded image {img_index} on page {page_num} of '{original_filename}'.")
                    continue

                if img_ext.lower() in ["png", "jpeg", "jpg", "gif", "bmp"]:
                    try:
                        image = Image.open(io.BytesIO(img_bytes))
                        image_np = np.array(image)
                        
                        ocr_task_future = loop.run_in_executor(
                            ocr_llm_executor,
                            reader.readtext,
                            image_np
                        )
                        ocr_tasks.append(ocr_task_future)
                        all_extracted_text_segments.append(f"--- Page {page_num + 1}, Embedded Image {img_index + 1} (OCR Scheduled) ---\n")
                    except Exception as img_e:
                        logger.error(f"Failed to prepare embedded image {img_index} on page {page_num} of '{original_filename}' for OCR: {img_e}", exc_info=True)
                else:
                    logger.warning(f"Skipping unsupported embedded image format '{img_ext}' on page {page_num}, image {img_index} of '{original_filename}'.")

            if perform_full_page_ocr:
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(1, 1)) # Render page to pixmap at 72dpi
                    img_bytes_from_page = pix.tobytes("png") # Convert pixmap to PNG bytes

                    if not img_bytes_from_page:
                        logger.warning(f"No bytes generated when rendering page {page_num} for full page OCR of '{original_filename}'.")
                        continue

                    image_from_page = Image.open(io.BytesIO(img_bytes_from_page))
                    image_from_page_np = np.array(image_from_page)

                    ocr_task_future = loop.run_in_executor(
                        ocr_llm_executor,
                        reader.readtext,
                        image_from_page_np
                    )
                    ocr_tasks.append(ocr_task_future)
                    all_extracted_text_segments.append(f"--- Page {page_num + 1} (Full Page OCR Scheduled) ---\n")
                except Exception as e_page_ocr:
                    logger.error(f"Could not render or prepare page {page_num} of '{original_filename}' for full page OCR: {e_page_ocr}", exc_info=True)
            else:
                logger.debug(f"Skipping full page OCR for page {page_num} of '{original_filename}' based on content heuristics.")

        doc.close()

        logger.info(f"Waiting for OCR tasks for '{original_filename}' to complete...")
        completed_ocr_results = await asyncio.gather(*ocr_tasks, return_exceptions=True)
        logger.info(f"All OCR tasks for '{original_filename}' completed.")

        for res in completed_ocr_results:
            if isinstance(res, Exception):
                logger.error(f"An OCR task failed for '{original_filename}': {res}", exc_info=True)
            else:
                extracted_text_from_ocr = [text for (bbox, text, prob) in res]
                if extracted_text_from_ocr:
                    combined_ocr_text_segment = "\n".join(extracted_text_from_ocr)
                    all_extracted_text_segments.append(combined_ocr_text_segment)

    elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        logger.info(f"Starting DOCX text extraction for '{original_filename}'.")
        # Run synchronous DOCX extraction in a thread pool
        all_extracted_text_segments.append(await asyncio.to_thread(_extract_text_from_docx, file_bytes))
        logger.info(f"Text extracted from DOCX '{original_filename}'.")
        
    elif content_type == "application/msword":
        logger.info(f"Starting DOC text extraction for '{original_filename}'.")
        # Run synchronous DOC extraction in a thread pool
        all_extracted_text_segments.append(await asyncio.to_thread(_extract_text_from_doc, file_bytes))
        logger.info(f"Text extracted from DOC '{original_filename}'.")

    else:
        logger.error(f"Unsupported content type '{content_type}' for '{original_filename}'.")
        raise ValueError(f"Unsupported file type for processing: {content_type}")

    full_resume_text = "\n".join(all_extracted_text_segments)

    # Basic check for empty resume text
    if not full_resume_text.strip():
        logger.warning(f"Extracted no meaningful text from '{original_filename}'. Skipping LLM call.")
        return {
            "original_filename": original_filename,
            "content_type": content_type,
            "parsed_data": { # Return an empty/default structure
                "basic_details": {"name": "N/A", "email": "N/A", "phone": "N/A", "links": {}},
                "technical_expertise": [],
                "certifications": [],
                "professional_summary": "N/A (No text extracted)",
                "professional_experience": []
            },
            "llm_usage": {"prompt_tokens": 0, "completion_tokens": 0, "total_tokens": 0},
            "status": "FAILED_NO_TEXT",
            "error_message": "Could not extract text from file."
        }


    logger.info(f"Calling LLM for '{original_filename}'...")
    parsed_resume_data, usage_info = await _call_llm_for_resume_parsing(full_resume_text)
    logger.info(f"LLM parsing completed for '{original_filename}'.")

    # Store the original filename and the parsed data
    return {
        "original_filename": original_filename,
        "content_type": content_type,
        "parsed_data": parsed_resume_data,
        "llm_usage": usage_info,
        "status": "SUCCESS" # Indicate success here
    }


async def process_batch_of_resumes(job_id: str, file_info_for_batch: List[Tuple[str, str, str]], temp_dir: str):
    """
    Background task to process a batch of resume files.
    Updates job status and stores results in the JobManager.
    'file_info_for_batch' is a list of (temp_file_path, original_filename, content_type)
    'temp_dir' is the temporary directory where the uploaded files are stored.
    """
    # Retrieve job data to get the target_format
    job_data = job_manager.get_job_status(job_id)
    if not job_data:
        logger.error(f"Job '{job_id}' not found during batch processing initiation.")
        return # Exit if job data is somehow missing

    target_format = job_data.target_format # Get the desired output format

    # Ensure job status is set to PROCESSING
    job_manager.update_job_status(job_id, JobStatus.PROCESSING)
    results = [] # This list will store the Dict[str, Any] results for each file
    
    total_files = len(file_info_for_batch)
    processed_count = 0

    try:
        for file_idx, (file_path, original_filename, content_type) in enumerate(file_info_for_batch):
            logger.info(f"Job '{job_id}': Processing file {file_idx + 1}/{total_files}: '{original_filename}'.")
            
            file_bytes = None
            try:
                # Read the temporary file
                async with aiofiles.open(file_path, 'rb') as f:
                    file_bytes = await f.read()
            except Exception as e:
                logger.error(f"Job '{job_id}': Error reading temporary file '{file_path}': {e}", exc_info=True)
                results.append({
                    "original_filename": original_filename,
                    "status": "FAILED",
                    "error_message": f"Could not read temporary file: {e}",
                    "parsed_data": None,
                    "llm_usage": None
                })
                # Clean up temp file immediately (even if reading failed)
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.debug(f"Job '{job_id}': Cleaned up temporary file '{file_path}'.")
                
                processed_count += 1
                progress = int((processed_count / total_files) * 100)
                job_manager.update_job_status(job_id, JobStatus.PROCESSING, progress=progress)
                continue # Skip to the next file in the loop

            try:
                # Process the single resume file
                single_file_result = await process_single_resume_file(file_bytes, original_filename, content_type)
                results.append(single_file_result)
                logger.info(f"Job '{job_id}': Successfully processed '{original_filename}'.")
            except Exception as e:
                # Catch any exception from process_single_resume_file and store a structured error object
                logger.error(f"Job '{job_id}': Error processing '{original_filename}': {e}", exc_info=True)
                results.append({
                    "original_filename": original_filename,
                    "status": "FAILED",
                    "error_message": str(e),
                    "parsed_data": None,
                    "llm_usage": None
                })
            finally:
                # Clean up the individual temporary file after processing
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.debug(f"Job '{job_id}': Cleaned up temporary file '{file_path}'.")
                
                # Update progress after each file, regardless of success/failure
                processed_count += 1
                progress = int((processed_count / total_files) * 100)
                job_manager.update_job_status(job_id, JobStatus.PROCESSING, progress=progress)

        # Generate ZIP file based on target_format
        output_ext = "pdf" if target_format == OutputFormat.PDF else "docx"
        zip_file_name = f"parsed_resumes_{job_id}_{output_ext}s.zip"
        zip_file_path = os.path.join(DOWNLOAD_DIR, zip_file_name)
        
        try:
            with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                for result_item in results:
                    # Only add successfully parsed items to the ZIP
                    if result_item and result_item.get("parsed_data") and result_item.get("status") == "SUCCESS":
                        parsed_data = result_item["parsed_data"]
                        original_fname = result_item.get("original_filename", "unknown_file")
                        
                        extracted_name = parsed_data.get("basic_details", {}).get("name", "").strip()
                        
                        # Determine base filename for output (e.g., "JOHN_DOE_RESUME.pdf" or "JOHN_DOE_RESUME.docx")
                        if extracted_name and extracted_name != "N/A":
                            sanitized_name_upper = re.sub(r'[^\w\s-]', '', extracted_name).strip().replace(' ', '_').upper()
                            output_filename = f"{sanitized_name_upper}_RESUME.{output_ext}"
                        else:
                            # Fallback if name not found, use original filename parts
                            name_part = os.path.splitext(original_fname)[0]
                            sanitized_name_part = re.sub(r'[^\w\s-]', '', name_part).strip().replace(' ', '_').upper()
                            output_filename = f"PARSED_{sanitized_name_part}_RESUME.{output_ext}"
                        
                        # Generate file bytes in a separate thread based on target_format
                        try:
                            if target_format == OutputFormat.PDF:
                                generated_bytes = await asyncio.to_thread(generate_pdf_from_json, parsed_data)
                            else: # OutputFormat.DOCX
                                generated_bytes = await asyncio.to_thread(generate_docx_from_json, parsed_data)
                            
                            # Handle potential duplicate filenames in zip to prevent overwrites
                            counter = 1
                            final_output_filename = output_filename
                            while final_output_filename in zf.namelist():
                                base, ext = os.path.splitext(output_filename)
                                final_output_filename = f"{base}_{counter}{ext}"
                                counter += 1

                            zf.writestr(final_output_filename, generated_bytes)
                            logger.info(f"Added '{final_output_filename}' to zip for job '{job_id}'.")
                        except Exception as gen_e:
                            logger.error(f"Failed to generate {output_ext.upper()} for '{original_fname}' (Job '{job_id}'): {gen_e}", exc_info=True)
                            # Add a placeholder text file to the zip indicating generation failure
                            zf.writestr(f"FAILED_{output_ext.upper()}_{os.path.splitext(original_fname)[0]}.txt", 
                                        f"{output_ext.upper()} generation failed for '{original_fname}': {gen_e}")
                    else:
                        # Add a text file indicating the failure to the zip for clarity if parsing failed
                        original_fname_for_log = result_item.get('original_filename', 'N/A') if result_item else 'N/A (result_item was None)'
                        error_msg_for_log = result_item.get('error_message', 'No specific error message') if result_item else 'N/A'
                        logger.warning(f"Skipping failed or invalid result for '{original_fname_for_log}' (Error: {error_msg_for_log}) when creating ZIP for job '{job_id}'.")
                        zf.writestr(f"FAILED_PARSE_{os.path.splitext(original_fname_for_log)[0]}.txt", 
                                    f"Parsing failed for '{original_fname_for_log}': {error_msg_for_log}. No {output_ext.upper()} generated.")

            job_manager.update_job_status(
                job_id, 
                JobStatus.COMPLETED, 
                results=results, 
                zip_file_path=zip_file_path, 
                progress=100
            )
            logger.info(f"Job '{job_id}' completed. ZIP file created at: {zip_file_path}")
        
        except Exception as zip_e:
            logger.error(f"Job '{job_id}': Failed to create ZIP file: {zip_e}", exc_info=True)
            job_manager.update_job_status(
                job_id, 
                JobStatus.FAILED, 
                error_message=f"Failed to create ZIP: {zip_e}", 
                results=results
            )
            # Ensure the potentially partially created zip is removed if an error occurred
            if os.path.exists(zip_file_path):
                os.remove(zip_file_path)

    except Exception as e:
        logger.error(f"Job '{job_id}': An unexpected error occurred during batch processing: {e}", exc_info=True)
        job_manager.update_job_status(job_id, JobStatus.FAILED, error_message=f"Batch processing failed: {e}")
    finally:
        # Crucial: Clean up the temporary upload directory for this job after processing
        # This will be called regardless of success or failure of the processing.
        # The actual temporary files within this directory should have been removed in the loop.
        # This just removes the empty directory itself.
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Job '{job_id}': Cleaned up temporary upload directory: {temp_dir}")
            except OSError as e:
                logger.error(f"Job '{job_id}': Error removing temporary directory {temp_dir}: {e}")